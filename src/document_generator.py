import io
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime

class DocumentGenerator:
    """
    Generates professional DOCX reports for Regulatory Intelligence findings.
    """
    
    def generate_regulatory_report_docx(self, df, query_term: str, stats_log: dict) -> io.BytesIO:
        doc = Document()
        
        # Title
        title = doc.add_heading(f"Regulatory Intelligence Report: {query_term}", 0)
        title.alignment = 1 # Center
        
        # Meta
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Scope: FDA, CPSC, UK MHRA, Health Canada")
        
        # Executive Summary / Stats
        doc.add_heading("Executive Summary", level=1)
        p = doc.add_paragraph()
        p.add_run("Search Sources and Yield:").bold = True
        
        for source, count in stats_log.items():
            doc.add_paragraph(f"- {source}: {count} records", style='List Bullet')
            
        doc.add_paragraph(f"Total Records Found: {len(df)}")
        
        # High Risk Section (if AI data exists)
        if "AI_Risk_Level" in df.columns:
            high_risk = df[df["AI_Risk_Level"] == "High"]
            if not high_risk.empty:
                h = doc.add_heading("⚠️ HIGH PRIORITY FINDINGS", level=1)
                run = h.runs[0]
                run.font.color.rgb = RGBColor(255, 0, 0)
                
                for idx, row in high_risk.iterrows():
                    self._add_record_to_doc(doc, row)
        
        # Full Findings
        doc.add_heading("Detailed Findings", level=1)
        
        # Sort by date if possible
        if 'Date' in df.columns:
            try:
                df = df.sort_values(by='Date', ascending=False)
            except: pass
            
        for idx, row in df.iterrows():
            # Skip high risk here if we want to avoid dups, or just list all
            self._add_record_to_doc(doc, row)
            
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _add_record_to_doc(self, doc, row):
        """Helper to format a single row in the doc."""
        title = row.get("Product", "Unknown Product")
        doc.add_heading(str(title), level=2)
        
        p = doc.add_paragraph()
        p.add_run("Date: ").bold = True
        p.add_run(str(row.get("Date", "N/A")))
        
        p = doc.add_paragraph()
        p.add_run("Source: ").bold = True
        p.add_run(str(row.get("Source", "N/A")))
        
        p = doc.add_paragraph()
        p.add_run("Reason: ").bold = True
        p.add_run(str(row.get("Reason", "N/A")))
        
        if "AI_Analysis" in row:
            p = doc.add_paragraph()
            p.add_run("AI Analysis: ").bold = True
            p.add_run(str(row.get("AI_Analysis", ""))).italic = True

        p = doc.add_paragraph()
        p.add_run("Link: ").bold = True
        p.add_run(str(row.get("Link", "")))
        
        doc.add_paragraph("_" * 50)
